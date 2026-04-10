#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Chunking script for Word documents using LangChain splitters.

Input modes:
1. Explicit list of files in input_files
2. All .docx files from input_dir

Chunking flow:
- Clean paragraphs
- Split by markdown headers with MarkdownHeaderTextSplitter
- Split oversized sections with RecursiveCharacterTextSplitter
- Build final chunks with min/max token limits
- Add adaptive overlap between neighboring chunks
"""

import re
import json
from pathlib import Path


CODE_BLOCK_RE = re.compile(r"```[\s\S]*?```")
DECORATIVE_LINE_RE = re.compile(r"^\s*[-*_`~]{3,}\s*$")


def get_token_counter():
    """Initialize token counter for GPT-style tokenization."""
    try:
        import tiktoken

        enc = tiktoken.get_encoding("cl100k_base")
        return lambda text: len(enc.encode(text))
    except ImportError:
        print("WARNING: tiktoken not installed. Using character-based counting as fallback.")
        print("Install tiktoken: pip install tiktoken")
        # Simple fallback: around 4 chars ~= 1 token
        return lambda text: len(text) // 4


def extract_document_title(paragraphs) -> str:
    """Extract document title from the first non-empty paragraph."""
    for para in paragraphs[:5]:
        text = para.text.strip()
        if text:
            return text.split("\n")[0][:100]
    return "Unknown Document"


def extract_sidebar_value(paragraphs, re_module) -> str:
    """Extract value from line matching 'sidebar: value'."""
    sidebar_re = re_module.compile(r"^\s*sidebar\s*:\s*(.+?)\s*$", re_module.IGNORECASE)
    for para in paragraphs:
        text = para.text.strip()
        if not text:
            continue
        match = sidebar_re.match(text)
        if match:
            value = match.group(1).strip()
            if value:
                return value
    return ""


def clean_paragraphs(paragraphs, re_module):
    """Return cleaned non-empty paragraphs excluding sidebar line."""
    cleaned = []
    sidebar_re = re_module.compile(r"^\s*sidebar\s*:\s*.+$", re_module.IGNORECASE)

    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if text and not sidebar_re.match(text):
            text = re_module.sub(r" +", " ", text)
            cleaned.append(text)

    return cleaned


def split_text_preserving_code_blocks(text):
    """Split text into segments while preserving whole fenced code blocks."""
    parts = []
    last_end = 0

    for match in CODE_BLOCK_RE.finditer(text):
        before = text[last_end:match.start()]
        if before.strip():
            parts.append((before, False))

        code_block = match.group(0)
        if code_block.strip():
            parts.append((code_block, True))

        last_end = match.end()

    tail = text[last_end:]
    if tail.strip():
        parts.append((tail, False))

    if not parts and text.strip():
        parts.append((text, False))

    return parts


def has_code_block(text):
    """Return True if text contains fenced code block."""
    return bool(CODE_BLOCK_RE.search(text))


def is_decorative_only_text(text):
    """Return True for blocks that contain only separator-like markdown lines."""
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return True
    return all(DECORATIVE_LINE_RE.match(line) for line in lines)


def count_chunk_tokens_excluding_metadata(text, token_counter, h1="", h2="", h3=""):
    """Count tokens excluding heading lines that mirror H1/H2/H3 metadata."""
    removable_lines = set()
    if h1:
        removable_lines.add(f"# {h1}".strip())
    if h2:
        removable_lines.add(f"## {h2}".strip())
    if h3:
        removable_lines.add(f"### {h3}".strip())

    if not removable_lines:
        return token_counter(text)

    filtered_lines = []
    for line in text.splitlines():
        normalized_line = line.strip()
        if normalized_line in removable_lines:
            continue
        filtered_lines.append(line)

    filtered_text = "\n".join(filtered_lines).strip()
    return token_counter(filtered_text)


def group_docs_by_h2(header_docs):
    """Group MarkdownHeaderTextSplitter docs into contiguous H2 sections."""
    groups = []
    current_key = None
    current_parts = []

    for doc in header_docs:
        text = doc.page_content.strip()
        if not text:
            continue

        metadata = doc.metadata or {}
        h1 = (metadata.get("h1") or "").strip()
        h2 = (metadata.get("h2") or "").strip()
        h3 = (metadata.get("h3") or "").strip()
        key = (h1, h2)

        part = {"text": text, "h3": h3}

        if current_key != key:
            if current_parts:
                groups.append(
                    {
                        "h1": current_key[0],
                        "h2": current_key[1],
                        "parts": current_parts,
                    }
                )
            current_key = key
            current_parts = [part]
        else:
            current_parts.append(part)

    if current_parts and current_key is not None:
        groups.append(
            {
                "h1": current_key[0],
                "h2": current_key[1],
                "parts": current_parts,
            }
        )

    return groups


def split_h2_group_into_h3_sections(h2_group):
    """Split one H2 group into contiguous H3 sections."""
    sections = []
    current_h3 = None
    current_parts = []

    for part in h2_group["parts"]:
        h3 = part["h3"]
        text = part["text"]

        if current_h3 is None:
            current_h3 = h3
            current_parts = [text]
            continue

        if h3 != current_h3:
            section_text = "\n\n".join(current_parts).strip()
            if section_text:
                sections.append({"h3": current_h3, "text": section_text})
            current_h3 = h3
            current_parts = [text]
        else:
            current_parts.append(text)

    if current_parts:
        section_text = "\n\n".join(current_parts).strip()
        if section_text:
            sections.append({"h3": current_h3, "text": section_text})

    return sections


def prepare_units_with_langchain(cleaned_paragraphs, token_counter, max_chunk_tokens):
    """Prepare units by hierarchy H2 -> H3, with no forced split for code chunks."""
    from langchain_text_splitters import (
        MarkdownHeaderTextSplitter,
        RecursiveCharacterTextSplitter,
    )

    markdown_text = "\n\n".join(cleaned_paragraphs).strip()
    if not markdown_text:
        return []

    header_splitter = MarkdownHeaderTextSplitter(
        headers_to_split_on=[
            ("#", "h1"),
            ("##", "h2"),
            ("###", "h3"),
        ],
        strip_headers=False,
    )

    recursive_splitter = RecursiveCharacterTextSplitter(
        chunk_size=max_chunk_tokens,
        chunk_overlap=0,
        length_function=token_counter,
        separators=["\n\n", "\n", ". ", "! ", "? ", "; ", " ", ""],
        keep_separator=True,
    )

    header_docs = header_splitter.split_text(markdown_text)
    h2_groups = group_docs_by_h2(header_docs)
    prepared_units = []

    for h2_group in h2_groups:
        h1_title = h2_group["h1"]
        h2_title = h2_group["h2"]
        h2_text = "\n\n".join(part["text"] for part in h2_group["parts"]).strip()
        if not h2_text or is_decorative_only_text(h2_text):
            continue

        if has_code_block(h2_text):
            prepared_units.append(
                {
                    "text": h2_text,
                    "has_code": True,
                    "h1": h1_title,
                    "h2": h2_title,
                    "h3": "",
                }
            )
            continue

        if token_counter(h2_text) <= max_chunk_tokens:
            prepared_units.append(
                {
                    "text": h2_text,
                    "has_code": False,
                    "h1": h1_title,
                    "h2": h2_title,
                    "h3": "",
                }
            )
            continue

        h3_sections = split_h2_group_into_h3_sections(h2_group)
        if len(h3_sections) <= 1:
            only_h3 = h3_sections[0]["h3"] if h3_sections else ""
            for piece in recursive_splitter.split_text(h2_text):
                text_piece = piece.strip()
                if text_piece:
                    prepared_units.append(
                        {
                            "text": text_piece,
                            "has_code": False,
                            "h1": h1_title,
                            "h2": h2_title,
                            "h3": only_h3,
                        }
                    )
            continue

        for h3_section in h3_sections:
            h3_title = h3_section["h3"]
            section_text = h3_section["text"]
            if not section_text or is_decorative_only_text(section_text):
                continue

            if has_code_block(section_text):
                prepared_units.append(
                    {
                        "text": section_text,
                        "has_code": True,
                        "h1": h1_title,
                        "h2": h2_title,
                        "h3": h3_title,
                    }
                )
                continue

            if token_counter(section_text) <= max_chunk_tokens:
                prepared_units.append(
                    {
                        "text": section_text,
                        "has_code": False,
                        "h1": h1_title,
                        "h2": h2_title,
                        "h3": h3_title,
                    }
                )
                continue

            for piece in recursive_splitter.split_text(section_text):
                text_piece = piece.strip()
                if text_piece and not is_decorative_only_text(text_piece):
                    prepared_units.append(
                        {
                            "text": text_piece,
                            "has_code": False,
                            "h1": h1_title,
                            "h2": h2_title,
                            "h3": h3_title,
                        }
                    )

    return prepared_units


def merge_small_chunks(chunk_info, token_counter, min_merge_tokens, max_chunk_tokens):
    """Merge too-small chunks with neighbors if merged text still fits the limit."""
    items = [
        {"text": text, "tokens": tokens, "has_code": has_code}
        for text, tokens, has_code in chunk_info
    ]
    index = 0

    while index < len(items):
        current = items[index]
        if current["tokens"] >= min_merge_tokens:
            index += 1
            continue

        merged = False

        if index + 1 < len(items):
            merged_text = current["text"] + "\n" + items[index + 1]["text"]
            merged_tokens = token_counter(merged_text)
            merged_has_code = current["has_code"] or items[index + 1]["has_code"]
            merge_limit = 1000 if merged_has_code else max_chunk_tokens
            if merged_tokens <= merge_limit:
                items[index + 1] = {
                    "text": merged_text,
                    "tokens": merged_tokens,
                    "has_code": merged_has_code,
                }
                del items[index]
                merged = True

        if not merged and index > 0:
            merged_text = items[index - 1]["text"] + "\n" + current["text"]
            merged_tokens = token_counter(merged_text)
            merged_has_code = items[index - 1]["has_code"] or current["has_code"]
            merge_limit = 1000 if merged_has_code else max_chunk_tokens
            if merged_tokens <= merge_limit:
                items[index - 1] = {
                    "text": merged_text,
                    "tokens": merged_tokens,
                    "has_code": merged_has_code,
                }
                del items[index]
                index -= 1
                merged = True

        if not merged:
            index += 1

    return [(item["text"], item["tokens"], item["has_code"]) for item in items]


def build_chunks(
    prepared_units,
    token_counter,
    min_chunk_tokens,
    max_chunk_tokens,
    max_chunk_tokens_with_overlap,
    min_merge_tokens,
    max_postprocess_merge_tokens,
):
    """Apply adaptive overlap to prepared units while preserving H1/H2/H3 metadata."""
    chunk_info = []
    for unit in prepared_units:
        text = unit["text"]
        if is_decorative_only_text(text):
            continue
        chunk_info.append(
            {
                "text": text,
                "tokens": count_chunk_tokens_excluding_metadata(
                    text,
                    token_counter,
                    unit.get("h1", ""),
                    unit.get("h2", ""),
                    unit.get("h3", ""),
                ),
                "has_code": unit["has_code"],
                "h1": unit.get("h1", ""),
                "h2": unit.get("h2", ""),
                "h3": unit.get("h3", ""),
            }
        )

    chunks_with_overlap = []

    for i, current in enumerate(chunk_info):
        chunk_text = current["text"]
        tokens = current["tokens"]
        chunk_has_code = current["has_code"]

        if i > 0:
            prev_tokens = chunk_info[i - 1]["tokens"]
            overlap_percent = 0.40 if prev_tokens < min_chunk_tokens else 0.15
            ideal_overlap_tokens = int(prev_tokens * overlap_percent)
            if chunk_has_code:
                overlap_tokens = ideal_overlap_tokens
            else:
                max_possible_overlap = max(0, max_chunk_tokens_with_overlap - tokens)
                overlap_tokens = min(ideal_overlap_tokens, max_possible_overlap)

            prev_lines = chunk_info[i - 1]["text"].split("\n")
            overlap_text = ""
            overlap_count = 0

            for line in reversed(prev_lines):
                line_tokens = token_counter(line)
                if overlap_count + line_tokens <= overlap_tokens:
                    overlap_text = line + "\n" + overlap_text
                    overlap_count += line_tokens
                else:
                    break

            if overlap_text:
                chunk_text = overlap_text + chunk_text

        chunks_with_overlap.append(
            {
                "text": chunk_text,
                "tokens": count_chunk_tokens_excluding_metadata(
                    chunk_text,
                    token_counter,
                    current.get("h1", ""),
                    current.get("h2", ""),
                    current.get("h3", ""),
                ),
                "has_code": chunk_has_code,
                "h1": current.get("h1", ""),
                "h2": current.get("h2", ""),
                "h3": current.get("h3", ""),
            }
        )

    return merge_small_chunks_postprocess(
        chunks_with_overlap,
        token_counter,
        min_merge_tokens,
        max_postprocess_merge_tokens,
    )


def merge_small_chunks_postprocess(chunks, token_counter, min_merge_tokens, max_merge_tokens):
    """Merge too-small final chunks with neighbors in the same H1/H2/H3 scope."""
    items = [dict(chunk) for chunk in chunks]
    index = 0

    while index < len(items):
        current = items[index]
        if current["tokens"] >= min_merge_tokens:
            index += 1
            continue

        merged = False

        # Prefer merge with next chunk to keep forward narrative flow.
        if index + 1 < len(items):
            right = items[index + 1]
            same_scope = (
                current.get("h1", "") == right.get("h1", "")
                and current.get("h2", "") == right.get("h2", "")
                and current.get("h3", "") == right.get("h3", "")
            )
            if same_scope:
                merged_text = current["text"] + "\n" + right["text"]
                merged_tokens = count_chunk_tokens_excluding_metadata(
                    merged_text,
                    token_counter,
                    current.get("h1", ""),
                    current.get("h2", ""),
                    current.get("h3", ""),
                )
                if merged_tokens <= max_merge_tokens:
                    items[index + 1] = {
                        "text": merged_text,
                        "tokens": merged_tokens,
                        "has_code": current["has_code"] or right["has_code"],
                        "h1": current.get("h1", ""),
                        "h2": current.get("h2", ""),
                        "h3": current.get("h3", ""),
                    }
                    del items[index]
                    merged = True

        if not merged and index > 0:
            left = items[index - 1]
            same_scope = (
                current.get("h1", "") == left.get("h1", "")
                and current.get("h2", "") == left.get("h2", "")
                and current.get("h3", "") == left.get("h3", "")
            )
            if same_scope:
                merged_text = left["text"] + "\n" + current["text"]
                merged_tokens = count_chunk_tokens_excluding_metadata(
                    merged_text,
                    token_counter,
                    left.get("h1", ""),
                    left.get("h2", ""),
                    left.get("h3", ""),
                )
                if merged_tokens <= max_merge_tokens:
                    items[index - 1] = {
                        "text": merged_text,
                        "tokens": merged_tokens,
                        "has_code": left["has_code"] or current["has_code"],
                        "h1": left.get("h1", ""),
                        "h2": left.get("h2", ""),
                        "h3": left.get("h3", ""),
                    }
                    del items[index]
                    index -= 1
                    merged = True

        if not merged:
            index += 1

    return items


def collect_input_files(input_files, input_dir):
    """Return unique existing .docx files to process."""
    resolved_files = []

    if input_files:
        for file_path in input_files:
            path_obj = Path(file_path)
            if path_obj.exists() and path_obj.suffix.lower() == ".docx":
                resolved_files.append(path_obj)
            else:
                print(f"WARNING: skipped invalid file: {file_path}")

    if input_dir:
        dir_path = Path(input_dir)
        if not dir_path.exists():
            print(f"WARNING: input directory not found: {input_dir}")
        else:
            resolved_files.extend(sorted(dir_path.rglob("*.docx")))

    unique_files = []
    seen = set()
    for file_path in resolved_files:
        normalized = str(file_path.resolve()).lower()
        if normalized not in seen:
            seen.add(normalized)
            unique_files.append(file_path)

    return unique_files


def sanitize_filename_part(value):
    """Prepare filename part for Windows-safe filenames."""
    sanitized = re.sub(r'[<>:"/\\|?*]+', "_", value).strip()
    sanitized = re.sub(r"\s+", "_", sanitized)
    return sanitized or "unknown"


def build_output_filename(file_path):
    """Build output filename with source parent folder prefix."""
    parent_name = sanitize_filename_part(file_path.parent.name)
    stem_name = sanitize_filename_part(file_path.stem)
    return f"{parent_name}_{stem_name}_chunks_langchain.txt"


def save_chunks(
    output_path,
    document_label,
    source_file,
    chunks_with_overlap,
    token_counter,
    min_chunk_tokens,
    max_chunk_tokens,
    max_chunk_tokens_with_overlap,
):
    """Save chunks of one document into a dedicated txt file."""
    chunk_token_counts = [chunk["tokens"] for chunk in chunks_with_overlap]
    avg_tokens = sum(chunk_token_counts) / len(chunks_with_overlap) if chunks_with_overlap else 0

    with open(output_path, "w", encoding="utf-8") as file_obj:
        file_obj.write(f"Source file: {source_file}\n")
        file_obj.write(f"Document: {document_label}\n")
        file_obj.write(f"Total chunks: {len(chunks_with_overlap)}\n")
        file_obj.write(
            f"Chunk size before overlap: {min_chunk_tokens}-{max_chunk_tokens} tokens "
            "(hierarchy: H2 -> H3; code chunks are kept whole)\n"
        )
        file_obj.write(
            f"Overlap: 15% (or 40% if previous chunk < {min_chunk_tokens} tokens), "
            f"target max with overlap: {max_chunk_tokens_with_overlap} tokens\n"
        )
        file_obj.write("=" * 80 + "\n\n")

        file_obj.write(f"Statistics: Min={min(chunk_token_counts) if chunks_with_overlap else 0} tokens, ")
        file_obj.write(f"Max={max(chunk_token_counts) if chunks_with_overlap else 0} tokens, ")
        file_obj.write(f"Avg={avg_tokens:.0f} tokens\n")
        file_obj.write("=" * 80 + "\n\n")

        for index, chunk in enumerate(chunks_with_overlap, 1):
            chunk_tokens = chunk["tokens"]
            chunk_text = chunk["text"]
            file_obj.write(f"--- CHUNK {index} ({chunk_tokens} tokens / {len(chunk_text)} chars) ---\n")
            file_obj.write(f"Document: {document_label}\n\n")
            file_obj.write(f"H1: {chunk.get('h1', '')}\n")
            file_obj.write(f"H2: {chunk.get('h2', '')}\n\n")
            file_obj.write(f"H3: {chunk.get('h3', '')}\n\n")
            file_obj.write(f"{chunk_text}\n\n")


def save_chunks_jsonl(
    output_path,
    document_label,
    source_file,
    chunks_with_overlap,
    token_counter,
):
    """Save chunks to JSONL (one JSON object per chunk line)."""
    with open(output_path, "w", encoding="utf-8") as file_obj:
        for index, chunk in enumerate(chunks_with_overlap, 1):
            payload = {
                "chunk_index": index,
                "document": document_label,
                "h1": chunk.get("h1", ""),
                "h2": chunk.get("h2", ""),
                "h3": chunk.get("h3", ""),
                "source_file": source_file,
                "tokens": chunk["tokens"],
                "chars": len(chunk["text"]),
                "contains_code": chunk["has_code"],
                "text": chunk["text"],
            }
            file_obj.write(json.dumps(payload, ensure_ascii=False) + "\n")


def process_document(file_path, output_dir, token_counter, min_chunk_tokens, max_chunk_tokens, re_module, document_factory, skip_code_chunks=False):
    """Process one document and save chunking result."""
    print(f"Reading: {file_path}")
    doc = document_factory(file_path)

    doc_title = extract_document_title(doc.paragraphs)
    print(f"Document title: {doc_title}")

    sidebar_value = extract_sidebar_value(doc.paragraphs, re_module)
    document_label = sidebar_value if sidebar_value else doc_title
    if sidebar_value:
        print(f"Sidebar label: {sidebar_value}")

    cleaned_paragraphs = clean_paragraphs(doc.paragraphs, re_module)
    print(f"Extracted: {len(cleaned_paragraphs)} paragraphs")

    prepared_units = prepare_units_with_langchain(
        cleaned_paragraphs,
        token_counter,
        max_chunk_tokens,
    )
    if skip_code_chunks:
        before = len(prepared_units)
        prepared_units = [u for u in prepared_units if not u["has_code"]]
        print(f"Skipped {before - len(prepared_units)} code unit(s)")
    print(
        "Prepared: "
        f"{len(prepared_units)} units using MarkdownHeaderTextSplitter + RecursiveCharacterTextSplitter"
    )

    chunks_with_overlap = build_chunks(
        prepared_units,
        token_counter,
        min_chunk_tokens,
        max_chunk_tokens,
        max_chunk_tokens + 105,
        350,
        1000,
    )
    print(f"Created {len(chunks_with_overlap)} chunks with adaptive overlap")

    output_path = output_dir / build_output_filename(file_path)
    save_chunks(
        output_path,
        document_label,
        str(file_path),
        chunks_with_overlap,
        token_counter,
        min_chunk_tokens,
        max_chunk_tokens,
        max_chunk_tokens + 105,
    )

    jsonl_output_path = output_path.with_suffix(".jsonl")
    save_chunks_jsonl(
        jsonl_output_path,
        document_label,
        str(file_path),
        chunks_with_overlap,
        token_counter,
    )

    print(f"Saved to: {output_path}")
    print(f"Saved to: {jsonl_output_path}")
    if chunks_with_overlap:
        chunk_token_counts = [chunk["tokens"] for chunk in chunks_with_overlap]
        print(
            "Chunk tokens: "
            f"min={min(chunk_token_counts)}, "
            f"max={max(chunk_token_counts)}, "
            f"avg={sum(chunk_token_counts) / len(chunk_token_counts):.0f}"
        )


def main():
    import re

    try:
        from docx import Document
    except ImportError:
        print("ERROR: python-docx not installed. Install: pip install python-docx")
        return

    try:
        import langchain_text_splitters  # noqa: F401
    except ImportError:
        print("ERROR: langchain-text-splitters not installed. Install: pip install langchain-text-splitters")
        return

    token_counter = get_token_counter()

    min_chunk_tokens = 500
    max_chunk_tokens = 700

    # Установите True, чтобы пропустить все чанки, содержащие код-блоки
    skip_code_chunks = False

    input_files = [
        # r"c:\Users\ar.kartavtsev\Desktop\Chunking\Test2.docx",
        # r"c:\Users\ar.kartavtsev\Desktop\Chunking\Test3.docx",
    ]
    input_dir = r"c:\Users\ar.kartavtsev\Desktop\Chunking\my-git-repo\input_docs"
    output_dir = Path(r"c:\Users\ar.kartavtsev\Desktop\Chunking\my-git-repo\chunks_output")

    files_to_process = collect_input_files(input_files, input_dir)
    if not files_to_process:
        print("ERROR: no .docx files found for processing.")
        print("Add files to input_files or place .docx files into input_dir.")
        return

    output_dir.mkdir(parents=True, exist_ok=True)
    print(f"Found {len(files_to_process)} file(s) for processing")

    for file_path in files_to_process:
        process_document(
            file_path,
            output_dir,
            token_counter,
            min_chunk_tokens,
            max_chunk_tokens,
            re,
            Document,
            skip_code_chunks=skip_code_chunks,
        )
        print("-" * 80)


if __name__ == "__main__":
    main()
