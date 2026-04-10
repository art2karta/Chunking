#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для чанкинга текста из Word документов.

Поддерживает 2 режима ввода:
1. Список конкретных файлов в input_files
2. Все .docx из папки input_dir
"""

from pathlib import Path

def get_token_counter():
    """Инициализирует counter токенов для GPT-моделей"""
    try:
        import tiktoken
        enc = tiktoken.get_encoding("cl100k_base")
        return lambda text: len(enc.encode(text))
    except ImportError:
        print("WARNING: tiktoken not installed. Using character-based counting as fallback.")
        print("Install tiktoken: pip install tiktoken")
        # Простой fallback: примерно 4 символа = 1 токен
        return lambda text: len(text) // 4


def extract_document_title(paragraphs) -> str:
    """Извлекает название документа из первого непустого абзаца"""
    for para in paragraphs[:5]:  # Смотрим первые 5 абзацев
        text = para.text.strip()
        if text and len(text) > 0:
            # Берем первую строку или первые 100 символов
            title = text.split('\n')[0][:100]
            return title
    return "Unknown Document"


def extract_sidebar_value(paragraphs, re_module) -> str:
    """Извлекает значение из строки вида 'sidebar: value'."""
    sidebar_re = re_module.compile(r"^\s*sidebar\s*:\s*(.+?)\s*$", re_module.IGNORECASE)
    for para in paragraphs:
        text = para.text.strip()
        if not text:
            continue
        match = sidebar_re.match(text)
        if match:
            value = match.group(1).strip()
            if value:
                return value
    return ""


def clean_paragraphs(paragraphs, re_module):
    """Очищает и возвращает только непустые абзацы."""
    cleaned_paragraphs = []
    sidebar_re = re_module.compile(r"^\s*sidebar\s*:\s*.+$", re_module.IGNORECASE)
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if text and not sidebar_re.match(text):
            text = re_module.sub(r" +", " ", text)
            cleaned_paragraphs.append(text)
    return cleaned_paragraphs


def split_by_markdown_heading_level(paragraphs, level, re_module):
    """Делит список абзацев по markdown-заголовкам заданного уровня (#, ##, ###)."""
    heading_re = re_module.compile(rf"^#{{{level}}}\\s+")
    sections = []
    current_section = []

    for paragraph in paragraphs:
        if heading_re.match(paragraph) and current_section:
            sections.append(current_section)
            current_section = [paragraph]
        else:
            current_section.append(paragraph)

    if current_section:
        sections.append(current_section)

    return sections


def split_paragraphs_to_sentences(paragraphs, re_module):
    """Разбивает абзацы на предложения для более мелкого уровня деления."""
    sentence_re = re_module.compile(r"(?<=[.!?])\s+")
    sentences = []

    for paragraph in paragraphs:
        parts = [part.strip() for part in sentence_re.split(paragraph) if part.strip()]
        if parts:
            sentences.extend(parts)

    return sentences


def split_text_by_token_limit(text, token_counter, max_chunk_tokens):
    """Делит длинный текст по символам так, чтобы каждый кусок укладывался в лимит токенов."""
    chunks = []
    start = 0
    text_len = len(text)

    while start < text_len:
        lo = start + 1
        hi = text_len
        best = start + 1

        while lo <= hi:
            mid = (lo + hi) // 2
            candidate = text[start:mid]
            if token_counter(candidate) <= max_chunk_tokens:
                best = mid
                lo = mid + 1
            else:
                hi = mid - 1

        piece = text[start:best].strip()
        if piece:
            chunks.append(piece)

        start = best

    return chunks


def recursive_hierarchical_split(paragraphs, token_counter, max_chunk_tokens, re_module, stage_index=0):
    """
    Каскадно делит текст по уровням:
    H2 -> H3 -> абзацы -> предложения -> символы.
    """
    if not paragraphs:
        return []

    text = "\n".join(paragraphs)
    if token_counter(text) <= max_chunk_tokens:
        return [text]

    stages = ["h2", "h3", "paragraph", "sentence", "char"]
    if stage_index >= len(stages):
        return [text]

    stage = stages[stage_index]

    if stage == "h2":
        sections = split_by_markdown_heading_level(paragraphs, 2, re_module)
        if len(sections) == 1:
            return recursive_hierarchical_split(
                paragraphs,
                token_counter,
                max_chunk_tokens,
                re_module,
                stage_index + 1,
            )

        result = []
        for section in sections:
            result.extend(
                recursive_hierarchical_split(
                    section,
                    token_counter,
                    max_chunk_tokens,
                    re_module,
                    stage_index + 1,
                )
            )
        return result

    if stage == "h3":
        sections = split_by_markdown_heading_level(paragraphs, 3, re_module)
        if len(sections) == 1:
            return recursive_hierarchical_split(
                paragraphs,
                token_counter,
                max_chunk_tokens,
                re_module,
                stage_index + 1,
            )

        result = []
        for section in sections:
            result.extend(
                recursive_hierarchical_split(
                    section,
                    token_counter,
                    max_chunk_tokens,
                    re_module,
                    stage_index + 1,
                )
            )
        return result

    if stage == "paragraph":
        result = []
        for paragraph in paragraphs:
            result.extend(
                recursive_hierarchical_split(
                    [paragraph],
                    token_counter,
                    max_chunk_tokens,
                    re_module,
                    stage_index + 1,
                )
            )
        return result

    if stage == "sentence":
        sentences = split_paragraphs_to_sentences(paragraphs, re_module)
        if not sentences:
            return recursive_hierarchical_split(
                paragraphs,
                token_counter,
                max_chunk_tokens,
                re_module,
                stage_index + 1,
            )

        result = []
        for sentence in sentences:
            result.extend(
                recursive_hierarchical_split(
                    [sentence],
                    token_counter,
                    max_chunk_tokens,
                    re_module,
                    stage_index + 1,
                )
            )
        return result

    if stage == "char":
        return split_text_by_token_limit(text, token_counter, max_chunk_tokens)

    return [text]


def prepare_units_for_chunking(cleaned_paragraphs, token_counter, max_chunk_tokens, re_module):
    """Готовит единицы текста через иерархическое деление до сборки чанков."""
    return recursive_hierarchical_split(
        cleaned_paragraphs,
        token_counter,
        max_chunk_tokens,
        re_module,
    )


def build_chunks(cleaned_paragraphs, token_counter, min_chunk_tokens, max_chunk_tokens):
    """Строит чанки с адаптивным перекрытием."""
    chunks = []
    current_chunk = []
    current_tokens = 0
    chunk_info = []

    for para in cleaned_paragraphs:
        para_tokens = token_counter(para)

        if current_tokens + para_tokens > max_chunk_tokens:
            if current_chunk and current_tokens >= min_chunk_tokens:
                chunk_text = "\n".join(current_chunk)
                chunks.append(chunk_text)
                chunk_info.append((chunk_text, current_tokens))
                current_chunk = [para]
                current_tokens = para_tokens
            elif (not current_chunk) or (current_tokens < min_chunk_tokens and para_tokens <= 200):
                current_chunk.append(para)
                current_tokens += para_tokens
            elif current_chunk:
                chunk_text = "\n".join(current_chunk)
                chunks.append(chunk_text)
                chunk_info.append((chunk_text, current_tokens))
                current_chunk = [para]
                current_tokens = para_tokens
        else:
            current_chunk.append(para)
            current_tokens += para_tokens

    if current_chunk:
        chunk_text = "\n".join(current_chunk)
        chunks.append(chunk_text)
        chunk_info.append((chunk_text, current_tokens))

    chunks_with_overlap = []
    for i, (chunk_text, tokens) in enumerate(chunk_info):
        if i > 0:
            prev_tokens = chunk_info[i - 1][1]
            overlap_percent = 0.40 if prev_tokens < min_chunk_tokens else 0.15
            ideal_overlap_tokens = int(prev_tokens * overlap_percent)
            max_possible_overlap = max(0, max_chunk_tokens - tokens)
            overlap_tokens = min(ideal_overlap_tokens, max_possible_overlap)

            prev_chunk_paragraphs = chunk_info[i - 1][0].split("\n")
            overlap_text = ""
            overlap_count = 0

            for para in reversed(prev_chunk_paragraphs):
                para_tokens = token_counter(para)
                if overlap_count + para_tokens <= overlap_tokens:
                    overlap_text = para + "\n" + overlap_text
                    overlap_count += para_tokens
                else:
                    break

            if overlap_text:
                chunk_text = overlap_text + chunk_text

        chunks_with_overlap.append(chunk_text)

    return chunks_with_overlap


def collect_input_files(input_files, input_dir):
    """Возвращает список существующих .docx файлов для обработки."""
    resolved_files = []

    if input_files:
        for file_path in input_files:
            path_obj = Path(file_path)
            if path_obj.exists() and path_obj.suffix.lower() == ".docx":
                resolved_files.append(path_obj)
            else:
                print(f"WARNING: skipped invalid file: {file_path}")

    if input_dir:
        dir_path = Path(input_dir)
        if not dir_path.exists():
            print(f"WARNING: input directory not found: {input_dir}")
        else:
            resolved_files.extend(sorted(dir_path.rglob("*.docx")))

    unique_files = []
    seen = set()
    for file_path in resolved_files:
        normalized = str(file_path.resolve()).lower()
        if normalized not in seen:
            seen.add(normalized)
            unique_files.append(file_path)

    return unique_files


def save_chunks(output_path, document_label, source_file, chunks_with_overlap, token_counter, min_chunk_tokens, max_chunk_tokens):
    """Сохраняет чанки одного документа в отдельный txt файл."""
    chunk_token_counts = [token_counter(chunk) for chunk in chunks_with_overlap]
    avg_tokens = sum(chunk_token_counts) / len(chunks_with_overlap) if chunks_with_overlap else 0

    with open(output_path, "w", encoding="utf-8") as file_obj:
        file_obj.write(f"Source file: {source_file}\n")
        file_obj.write(f"Document: {document_label}\n")
        file_obj.write(f"Total chunks: {len(chunks_with_overlap)}\n")
        file_obj.write(
            f"Chunk size: {min_chunk_tokens}-{max_chunk_tokens} tokens "
            "(hierarchical split: H2 -> H3 -> paragraph -> sentence -> char)\n"
        )
        file_obj.write(
            f"Overlap: 15% (or 40% if previous chunk < {min_chunk_tokens} tokens), "
            f"but never exceeds {max_chunk_tokens} tokens total\n"
        )
        file_obj.write("=" * 80 + "\n\n")

        file_obj.write(f"Statistics: Min={min(chunk_token_counts) if chunks_with_overlap else 0} tokens, ")
        file_obj.write(f"Max={max(chunk_token_counts) if chunks_with_overlap else 0} tokens, ")
        file_obj.write(f"Avg={avg_tokens:.0f} tokens\n")
        file_obj.write("=" * 80 + "\n\n")

        for index, chunk in enumerate(chunks_with_overlap, 1):
            chunk_tokens = chunk_token_counts[index - 1]
            file_obj.write(f"--- CHUNK {index} ({chunk_tokens} tokens / {len(chunk)} chars) ---\n")
            file_obj.write(f"Document: {document_label}\n")
            file_obj.write("\n")
            file_obj.write(f"{chunk}\n\n")


def process_document(file_path, output_dir, token_counter, min_chunk_tokens, max_chunk_tokens, re_module, document_factory):
    """Обрабатывает один документ и сохраняет результат."""
    print(f"Reading: {file_path}")
    doc = document_factory(file_path)

    doc_title = extract_document_title(doc.paragraphs)
    print(f"Document title: {doc_title}")

    sidebar_value = extract_sidebar_value(doc.paragraphs, re_module)
    document_label = sidebar_value if sidebar_value else doc_title
    if sidebar_value:
        print(f"Sidebar label: {sidebar_value}")

    cleaned_paragraphs = clean_paragraphs(doc.paragraphs, re_module)
    print(f"Extracted: {len(cleaned_paragraphs)} paragraphs")

    prepared_units = prepare_units_for_chunking(
        cleaned_paragraphs,
        token_counter,
        max_chunk_tokens,
        re_module,
    )
    print(f"Prepared: {len(prepared_units)} hierarchical units (H2 -> H3 -> paragraph -> sentence -> char)")

    chunks_with_overlap = build_chunks(
        prepared_units,
        token_counter,
        min_chunk_tokens,
        max_chunk_tokens,
    )
    print(f"Created {len(chunks_with_overlap)} chunks with adaptive overlap")

    output_path = output_dir / f"{file_path.stem}_chunks.txt"
    save_chunks(
        output_path,
        document_label,
        str(file_path),
        chunks_with_overlap,
        token_counter,
        min_chunk_tokens,
        max_chunk_tokens,
    )

    print(f"Saved to: {output_path}")
    if chunks_with_overlap:
        chunk_token_counts = [token_counter(chunk) for chunk in chunks_with_overlap]
        print(
            "Chunk tokens: "
            f"min={min(chunk_token_counts)}, "
            f"max={max(chunk_token_counts)}, "
            f"avg={sum(chunk_token_counts) / len(chunk_token_counts):.0f}"
        )


def main():
    import re
    
    try:
        from docx import Document
    except ImportError:
        print("ERROR: python-docx not installed. Install: pip install python-docx")
        return
    
    # Инициализируем counter токенов
    token_counter = get_token_counter()
    
    # Параметры
    min_chunk_tokens = 600
    max_chunk_tokens = 800
    input_files = [
        # r"c:\Users\ar.kartavtsev\Desktop\Chunking\Test2.docx",
        # r"c:\Users\ar.kartavtsev\Desktop\Chunking\Test3.docx",
    ]
    input_dir = r"c:\Users\ar.kartavtsev\Desktop\Chunking\my-git-repo\input_docs"
    output_dir = Path(r"c:\Users\ar.kartavtsev\Desktop\Chunking\my-git-repo\chunks_output")

    files_to_process = collect_input_files(input_files, input_dir)
    if not files_to_process:
        print("ERROR: no .docx files found for processing.")
        print("Add files to input_files or place .docx files into input_dir.")
        return

    output_dir.mkdir(parents=True, exist_ok=True)
    print(f"Found {len(files_to_process)} file(s) for processing")

    for file_path in files_to_process:
        process_document(
            file_path,
            output_dir,
            token_counter,
            min_chunk_tokens,
            max_chunk_tokens,
            re,
            Document,
        )
        print("-" * 80)

if __name__ == "__main__":
    main()
